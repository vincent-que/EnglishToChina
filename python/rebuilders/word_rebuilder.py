"""Word Rebuilder — replaces translated text in DOCX while preserving all formatting.

Key strategy: "run-level style preservation"
- For paragraphs: find the dominant run (most text), place all translated text there,
  clear other runs. This preserves font family, size, bold, italic, color, etc.
- For table cells: same dominant-run strategy per cell.
- Merged cells: tracked via original model, skip continuation cells.

This avoids the naive `cell.text = new_text` which destroys all formatting.
"""

import os
import sys
import logging
import re
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

logger = logging.getLogger(__name__)


def _match_key(text: str) -> str:
    return re.sub(r'\s+', ' ', text or '').strip()


class WordRebuilder:
    def rebuild(self, doc_model: dict, output_path: str,
                source_path: str = None) -> dict:
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)

        source_path = self._resolve_template_path(doc_model, source_path)
        translations = doc_model.get('translations', {})
        if not translations:
            logger.warning('没有翻译结果，直接复制原文')
            if source_path and os.path.exists(source_path):
                import shutil
                shutil.copy2(source_path, output_path)
            return {'outputPath': output_path}

        if (source_path and os.path.exists(source_path)
                and source_path.lower().endswith(('.docx', '.doc'))):
            logger.info('从源文件重建: %s', source_path)
            doc = Document(source_path)
            self._replace_in_place(doc, translations, doc_model)
            doc.save(output_path)
        else:
            logger.info('从模型新建文档')
            doc = self._build_from_model(doc_model, translations)
            doc.save(output_path)

        exists = os.path.exists(output_path)
        size = os.path.getsize(output_path) if exists else 0
        logger.info('保存完成: %s (%.1f KB)', output_path, size / 1024)
        return {'outputPath': output_path}

    @staticmethod
    def _resolve_template_path(doc_model: dict, source_path: str = None) -> str:
        converted_docx = doc_model.get('meta', {}).get('convertedDocx', '')
        if converted_docx and os.path.exists(converted_docx):
            return converted_docx
        return source_path

    # ── In-place replacement (preserves full formatting) ─────────────

    def _replace_in_place(self, doc: Document, translations: dict,
                          doc_model: dict):
        """Walk the document body in order, replacing paragraph and table text."""
        # Build lookup: originalText → translated
        block_map = {}
        for page in doc_model.get('pages', []):
            for block in page.get('blocks', []):
                block_id = block.get('id', '')
                translated = translations.get(block_id, '')
                if translated:
                    original = block.get('originalText', block.get('text', ''))
                    key = _match_key(original)
                    if key:
                        block_map[key] = translated

        body = doc.element.body
        table_counter = 0

        for child in body:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

            if tag == 'p':
                full_text = self._get_xml_text(child).strip()
                if not full_text:
                    continue

                translated = block_map.get(_match_key(full_text), '')
                if translated:
                    self._set_xml_para_text(child, translated)

            elif tag == 'tbl':
                if table_counter < len(doc.tables):
                    table_model = self._find_table_model(doc_model, table_counter)
                    if table_model:
                        doc_table = doc.tables[table_counter]
                        self._replace_table_cells(doc_table, translations, table_model)
                    table_counter += 1

    def _replace_table_cells(self, doc_table, translations: dict,
                             table_model: dict):
        """Replace text in table cells using first-run replacement.

        Uses cell._tc identity to deduplicate merged cells — python-docx's
        row.cells expands merged cells so the same _tc appears multiple times.
        """
        table_id = table_model.get('id', '')
        text_map = self._build_table_text_map(table_model, translations)
        processed_tcs = set()  # Track underlying XML elements to deduplicate

        for row_idx, row_model in enumerate(table_model.get('cells', [])):
            for col_idx, cell_model in enumerate(row_model):
                # Skip continuation cells from merge
                if cell_model.get('colSpan', 1) == 0:
                    continue

                cell_key = f"{table_id}_r{row_idx}_c{col_idx}"
                translated = translations.get(cell_key, '')
                original = cell_model.get('originalText', cell_model.get('text', '')).strip()
                if not original:
                    continue
                translated = translated or text_map.get(_match_key(original), '')
                if not translated:
                    continue

                if row_idx < len(doc_table.rows) and col_idx < len(doc_table.columns):
                    try:
                        doc_cell = doc_table.cell(row_idx, col_idx)
                        # Core fix: skip if this merged cell was already written
                        if doc_cell._tc in processed_tcs:
                            continue
                        processed_tcs.add(doc_cell._tc)
                        self._set_cell_text(doc_cell, translated)
                    except Exception as e:
                        logger.warning('表格单元格写入失败 [%d,%d]: %s',
                                       row_idx, col_idx, e)

        self._replace_unmatched_table_cells(doc_table, text_map, processed_tcs)

    def _build_table_text_map(self, table_model: dict, translations: dict) -> dict:
        table_id = table_model.get('id', '')
        text_map = {}
        for row_idx, row in enumerate(table_model.get('cells', [])):
            for col_idx, cell in enumerate(row):
                original = cell.get('originalText', cell.get('text', '')).strip()
                if not original:
                    continue
                cell_key = f"{table_id}_r{row_idx}_c{col_idx}"
                translated = translations.get(cell_key, '')
                key = _match_key(original)
                if key and translated:
                    text_map[key] = translated
        return text_map

    def _replace_unmatched_table_cells(self, doc_table, text_map: dict, processed_tcs: set):
        for row in doc_table.rows:
            for cell in row.cells:
                if cell._tc in processed_tcs:
                    continue
                original = cell.text.strip()
                translated = text_map.get(_match_key(original), '')
                if translated:
                    processed_tcs.add(cell._tc)
                    self._set_cell_text(cell, translated)

    def _set_cell_text(self, cell, new_text: str):
        """Replace cell text preserving formatting using first-run replacement.

        Strategy:
        1. Write translated text to the FIRST run (inherits its font/size/bold/color)
        2. Clear all remaining runs' text
        3. Clear extra paragraphs in the cell
        """
        if not cell.paragraphs:
            return

        para = cell.paragraphs[0]
        self._set_run_text(para, new_text)

        # Clear remaining paragraphs in the cell
        for p in cell.paragraphs[1:]:
            for run in p.runs:
                run.text = ''

    def _set_run_text(self, para, new_text: str):
        """Replace paragraph text using first-run replacement.

        First-run strategy:
        - paragraph.runs[0].text = translated_text  (preserves font, size, bold, color)
        - paragraph.runs[1:].text = ""  (clear rest, keep XML structure)

        This is more reliable than "dominant-run" because the first run
        carries the paragraph-level formatting intent.
        """
        if not para.runs:
            # No runs — use the simple text setter (creates a single run)
            para.text = new_text
            return

        # Core: write to first run, clear all others
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ''

    # ── XML-level helpers ────────────────────────────────────────────

    def _get_xml_text(self, para_element) -> str:
        """Get full text from a paragraph XML element (all runs combined)."""
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        texts = []
        for r_elem in para_element.findall('.//w:r', ns):
            for t_elem in r_elem.findall('w:t', ns):
                if t_elem.text:
                    texts.append(t_elem.text)
        return ''.join(texts)

    def _set_xml_para_text(self, para_element, new_text: str):
        """Replace paragraph text at XML level using first-run replacement.

        This operates directly on the XML tree, which is necessary for
        paragraphs found via body iteration (not python-docx objects).

        First-run strategy: write translated text into the first <w:r>,
        clear all subsequent <w:r> text. Preserves the first run's
        <w:rPr> (font, size, bold, color, etc.).
        """
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        runs = para_element.findall('.//w:r', ns)
        if not runs:
            return

        # Core: first-run replacement
        first_run = runs[0]

        # Handle first run: replace t elements with new text
        first_t_elems = first_run.findall('w:t', ns)
        for t_elem in first_t_elems:
            first_run.remove(t_elem)
        new_t = etree.Element(qn('w:t'))
        new_t.text = new_text
        new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        first_run.append(new_t)

        # Clear all other runs' text
        for r_elem in runs[1:]:
            for t_elem in r_elem.findall('w:t', ns):
                t_elem.text = ''

    def _find_table_model(self, doc_model: dict, table_idx: int):
        """Find the table model by sequential index across all pages."""
        counter = 0
        for page in doc_model.get('pages', []):
            for table in page.get('tables', []):
                if counter == table_idx:
                    return table
                counter += 1
        return None

    # ── Build from scratch (fallback) ────────────────────────────────

    def _build_from_model(self, doc_model: dict, translations: dict) -> Document:
        """Fallback: build a new document from scratch (no source file)."""
        doc = Document()
        for page in doc_model.get('pages', []):
            for block in page.get('blocks', []):
                block_id = block.get('id', '')
                text = translations.get(block_id, block.get('text', ''))
                block_type = block.get('type', 'paragraph')

                if block_type == 'heading':
                    doc.add_heading(text, level=1)
                elif block_type == 'list':
                    doc.add_paragraph(text, style='List Bullet')
                else:
                    doc.add_paragraph(text)

            for table in page.get('tables', []):
                rows = table.get('rows', 0)
                cols = table.get('cols', 0)
                if rows == 0 or cols == 0:
                    continue

                doc_table = doc.add_table(rows=rows, cols=cols)
                doc_table.style = 'Table Grid'

                for row_idx, row in enumerate(table.get('cells', [])):
                    for col_idx, cell_model in enumerate(row):
                        if cell_model.get('colSpan', 1) == 0:
                            continue
                        if row_idx < rows and col_idx < cols:
                            cell_key = f"{table['id']}_r{row_idx}_c{col_idx}"
                            translated = translations.get(cell_key, cell_model.get('text', ''))
                            doc_table.cell(row_idx, col_idx).text = translated
        return doc
