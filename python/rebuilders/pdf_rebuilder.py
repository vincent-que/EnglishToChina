"""PDF Rebuilder — converts translated DOCX to PDF.

Strategy (in priority order):
1. docx2pdf — requires MS Word installed (best quality)
2. LibreOffice — cross-platform, headless conversion
3. Pure Python fallback — python-docx + ReportLab (no external deps)

The pure Python fallback reads the translated .docx, extracts paragraphs
and tables with formatting, and renders to PDF using ReportLab with CJK fonts.
"""

import os
import sys
import logging
import subprocess
import platform
import shutil
import tempfile

logger = logging.getLogger(__name__)

def get_cjk_font_candidates():
    """Return bundled and system CJK fonts in lookup order."""
    module_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..'))
    resource_root = os.environ.get(
        'ENGLISH_TO_CHINA_RESOURCES',
        os.path.join(module_root, 'resources'),
    )
    return [
        ('NotoSansSC', os.path.join(resource_root, 'fonts', 'NotoSansSC-Regular.otf')),
        ('SourceHanSansSC', os.path.join(resource_root, 'fonts', 'SourceHanSansSC-Regular.otf')),
        ('Microsoft YaHei', r'C:\Windows\Fonts\msyh.ttc'),
        ('SimHei', r'C:\Windows\Fonts\simhei.ttf'),
        ('SimSun', r'C:\Windows\Fonts\simsun.ttc'),
    ]


class PDFRebuilder:
    def __init__(self):
        self._cjk_font_name = None

    def rebuild(self, doc_model: dict, output_path: str,
                source_path: str = None, translated_docx_path: str = None) -> dict:
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)

        source_path = self._resolve_template_path(doc_model, source_path)
        # If no translated DOCX provided, build it first
        if not translated_docx_path or not os.path.exists(translated_docx_path):
            translated_docx_path = self._build_translated_docx(
                doc_model, output_path, source_path
            )

        # Try conversion methods in priority order
        converted = False
        if self._try_docx2pdf(translated_docx_path, output_path):
            converted = True
        elif self._try_libreoffice(translated_docx_path, output_path):
            converted = True
        elif self._try_pure_python(translated_docx_path, output_path):
            converted = True

        if not converted:
            raise RuntimeError(
                '无法转换 DOCX→PDF。请安装以下任一工具：\n'
                '- Microsoft Word（最佳兼容性）\n'
                '- LibreOffice（免费，apt install libreoffice-writer）\n'
                '- 或确保 resources/fonts/ 下有中文字体文件'
            )

        exists = os.path.exists(output_path)
        size = os.path.getsize(output_path) if exists else 0
        logger.info('PDF 重建完成: %s (%.1f KB)', output_path, size / 1024)
        return {'outputPath': output_path}

    @staticmethod
    def _resolve_template_path(doc_model: dict, source_path: str = None) -> str:
        converted_docx = doc_model.get('meta', {}).get('convertedDocx', '')
        if converted_docx and os.path.exists(converted_docx):
            return converted_docx
        return source_path

    def _build_translated_docx(self, doc_model: dict, output_path: str,
                               source_path: str) -> str:
        from rebuilders.word_rebuilder import WordRebuilder
        base = os.path.splitext(os.path.basename(output_path))[0]
        temp_dir = tempfile.mkdtemp(prefix='etc_pdf_rebuild_')
        docx_path = os.path.join(temp_dir, f'{base}_translated.docx')
        rebuilder = WordRebuilder()
        rebuilder.rebuild(doc_model, docx_path, source_path=source_path)
        logger.info('中间 DOCX 生成: %s', docx_path)
        return docx_path

    # ── Method 1: docx2pdf (requires MS Word) ────────────────────────

    def _try_docx2pdf(self, docx_path: str, pdf_path: str) -> bool:
        try:
            from docx2pdf import convert
            convert(docx_path, pdf_path)
            if os.path.exists(pdf_path):
                logger.info('DOCX→PDF 转换成功 (docx2pdf/MS Word)')
                return True
        except ImportError:
            logger.debug('docx2pdf 未安装')
        except Exception as e:
            logger.debug('docx2pdf 转换失败: %s', e)
        return False

    # ── Method 2: LibreOffice ────────────────────────────────────────

    def _try_libreoffice(self, docx_path: str, pdf_path: str) -> bool:
        lo_cmd = self._find_libreoffice()
        if not lo_cmd:
            logger.debug('LibreOffice 未找到')
            return False

        output_dir = os.path.dirname(pdf_path)
        try:
            result = subprocess.run(
                [lo_cmd, '--headless', '--convert-to', 'pdf',
                 '--outdir', output_dir, docx_path],
                capture_output=True, text=True, timeout=120,
            )
            if result.returncode == 0:
                base = os.path.splitext(os.path.basename(docx_path))[0]
                lo_output = os.path.join(output_dir, f'{base}.pdf')
                if os.path.exists(lo_output) and lo_output != pdf_path:
                    shutil.move(lo_output, pdf_path)
                if os.path.exists(pdf_path):
                    logger.info('DOCX→PDF 转换成功 (LibreOffice)')
                    return True
            else:
                logger.warning('LibreOffice 转换失败: %s', result.stderr[:200])
        except FileNotFoundError:
            logger.debug('LibreOffice 命令未找到')
        except subprocess.TimeoutExpired:
            logger.warning('LibreOffice 转换超时')
        except Exception as e:
            logger.debug('LibreOffice 转换异常: %s', e)
        return False

    @staticmethod
    def _find_libreoffice() -> str:
        system = platform.system()
        if system == 'Windows':
            candidates = [
                r'C:\Program Files\LibreOffice\program\soffice.exe',
                r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
            ]
            for p in candidates:
                if os.path.exists(p):
                    return p
        elif system == 'Darwin':
            lo_path = '/Applications/LibreOffice.app/Contents/MacOS/soffice'
            if os.path.exists(lo_path):
                return lo_path
        else:
            path = shutil.which('libreoffice') or shutil.which('soffice')
            if path:
                return path
        return shutil.which('soffice') or ''

    # ── Method 3: Pure Python (ReportLab) ────────────────────────────

    def _try_pure_python(self, docx_path: str, pdf_path: str) -> bool:
        """Render translated DOCX to PDF using python-docx + ReportLab."""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.units import mm
            from reportlab.pdfgen import canvas
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.pdfbase.cidfonts import UnicodeCIDFont
            from docx import Document
        except ImportError as e:
            logger.debug('ReportLab 鎴?python-docx 鏈畨瑁? %s', e)
            return False

        font_name = self._register_cjk_font(pdfmetrics, TTFont, UnicodeCIDFont)
        if not font_name:
            logger.warning('无法注册中文字体')
            return False

        try:
            doc = Document(docx_path)
        except Exception as e:
            logger.warning('无法打开 DOCX: %s', e)
            return False

        page_w, page_h = A4
        margin = 25 * mm
        c = canvas.Canvas(pdf_path, pagesize=A4)

        y = page_h - margin
        line_height = 14
        max_text_w = page_w - 2 * margin

        def new_page():
            nonlocal y
            c.showPage()
            y = page_h - margin

        def draw_wrapped(text, font_size, max_width=max_text_w, x=margin):
            nonlocal y
            if not text or not text.strip():
                return

            c.setFont(font_name, font_size)
            lines = self._wrap_text(text, font_name, font_size, max_width, c)

            for line in lines:
                if y < margin + line_height:
                    new_page()
                    c.setFont(font_name, font_size)
                c.drawString(x, y, line)
                y -= font_size * 1.5

        def paragraph_font_size(para):
            style_name = para.style.name.lower() if para.style else ''
            if 'heading 1' in style_name or 'title' in style_name:
                return 18, 6
            if 'heading 2' in style_name:
                return 15, 4
            if 'heading 3' in style_name:
                return 13, 3
            for run in para.runs:
                if run.font.size:
                    return run.font.size.pt, 0
            return 11, 0

        def draw_paragraph(para):
            nonlocal y
            text = para.text.strip()
            if not text:
                y -= line_height * 0.5
                if y < margin:
                    new_page()
                return

            font_size, extra_space = paragraph_font_size(para)
            y -= extra_space
            draw_wrapped(text, font_size)
            y -= 3

        def table_column_widths(table):
            if not table.rows or not table.columns:
                return []

            widths = []
            for cell in table.rows[0].cells:
                width = getattr(cell, 'width', None)
                widths.append(float(width.pt) if width else 0)

            if not any(widths):
                return [max_text_w / len(table.columns)] * len(table.columns)

            total = sum(widths)
            return [(width / total) * max_text_w for width in widths]

        def cell_col_span(cell):
            tc_pr = cell._tc.tcPr
            if tc_pr is None or tc_pr.gridSpan is None:
                return 1
            try:
                return max(1, int(tc_pr.gridSpan.val))
            except (TypeError, ValueError):
                return 1

        def draw_table(table):
            nonlocal y
            y -= 8
            if y < margin + 30:
                new_page()

            col_widths = table_column_widths(table)
            if not col_widths:
                return

            cell_font_size = 9
            for row in table.rows:
                max_lines = 1
                processed_tcs = set()
                for col_idx, cell in enumerate(row.cells):
                    if cell._tc in processed_tcs:
                        continue
                    processed_tcs.add(cell._tc)
                    span = cell_col_span(cell)
                    col_width = sum(
                        col_widths[min(idx, len(col_widths) - 1)]
                        for idx in range(col_idx, col_idx + span)
                    )
                    text = cell.text.strip()
                    if text:
                        lines = self._wrap_text(text, font_name, cell_font_size, col_width - 4, c)
                        max_lines = max(max_lines, len(lines))

                row_height = max_lines * cell_font_size * 1.5 + 6
                if y - row_height < margin:
                    new_page()

                x = margin
                processed_tcs = set()
                for col_idx, cell in enumerate(row.cells):
                    if cell._tc in processed_tcs:
                        continue
                    processed_tcs.add(cell._tc)
                    span = cell_col_span(cell)
                    col_width = sum(
                        col_widths[min(idx, len(col_widths) - 1)]
                        for idx in range(col_idx, col_idx + span)
                    )
                    c.rect(x, y - row_height, col_width, row_height)

                    text = cell.text.strip()
                    if text:
                        c.setFont(font_name, cell_font_size)
                        cell_lines = self._wrap_text(text, font_name, cell_font_size, col_width - 6, c)
                        text_y = y - cell_font_size * 1.2 - 3
                        for line in cell_lines:
                            if text_y < y - row_height + 2:
                                break
                            c.drawString(x + 3, text_y, line)
                            text_y -= cell_font_size * 1.3

                    x += col_width

                y -= row_height

            y -= 10

        paragraphs = iter(doc.paragraphs)
        tables = iter(doc.tables)
        for child in doc.element.body:
            tag = child.tag.rsplit('}', 1)[-1]
            if tag == 'p':
                draw_paragraph(next(paragraphs))
            elif tag == 'tbl':
                draw_table(next(tables))

        c.save()
        logger.info('纯 Python PDF 渲染完成: %s', pdf_path)
        return True

    def _register_cjk_font(self, pdfmetrics, TTFont, UnicodeCIDFont) -> str:
        """Register a CJK font for ReportLab. Returns font name or empty string."""
        for name, path in get_cjk_font_candidates():
            if not path or not os.path.exists(path):
                continue
            try:
                pdfmetrics.registerFont(TTFont(name, path))
                logger.info('注册字体: %s (%s)', name, path)
                return name
            except Exception as e:
                logger.debug('字体注册失败 %s: %s', name, e)

        # Fallback: CID font
        try:
            pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
            logger.info('注册 CID 字体: STSong-Light')
            return 'STSong-Light'
        except Exception:
            pass

        return ''

    @staticmethod
    def _wrap_text(text: str, font_name: str, font_size: float,
                   max_width: float, c) -> list:
        """Wrap text to fit within max_width."""
        if not text:
            return ['']

        lines = []
        current = ''
        for char in text:
            test = current + char
            try:
                w = c.stringWidth(test, font_name, font_size)
            except Exception:
                w = len(test) * font_size * 0.55
            if w > max_width and current:
                lines.append(current)
                current = char
            else:
                current = test
        if current:
            lines.append(current)
        return lines or ['']
