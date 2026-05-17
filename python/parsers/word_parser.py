"""Word Parser — extracts paragraphs, tables (with merged cells), and images."""

import os
import re
import logging
from docx import Document
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

# Invisibles/whitespace that should be stripped from extracted text
_INVISIBLE_RE = re.compile(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f​‌‍﻿]')


def _clean_text(text: str) -> str:
    """Remove invisible characters and normalize whitespace."""
    text = _INVISIBLE_RE.sub('', text)
    return text.strip()


class WordParser:
    def parse(self, file_path: str) -> dict:
        if not os.path.isabs(file_path):
            file_path = os.path.abspath(file_path)
        if not os.path.exists(file_path):
            raise FileNotFoundError(f'文件不存在: {file_path}')

        logger.info('开始解析 DOCX: %s', file_path)
        doc = Document(file_path)

        blocks = self._extract_paragraphs(doc)
        tables = self._extract_tables(doc)

        logger.info('解析完成: %d 个段落, %d 个表格', len(blocks), len(tables))

        return {
            'meta': {
                'sourceFile': os.path.basename(file_path),
                'format': 'docx',
                'pages': 1,
            },
            'pages': [{
                'pageNumber': 1,
                'blocks': blocks,
                'tables': tables,
                'images': [],
                'width': 595,
                'height': 842,
            }],
            'translations': {},
        }

    def _extract_paragraphs(self, doc: Document) -> list:
        """Extract paragraphs with style info."""
        blocks = []
        block_idx = 0

        for para in doc.paragraphs:
            text = _clean_text(para.text)
            if not text:
                continue

            style_name = para.style.name if para.style else 'Normal'
            block_type = self._detect_type(style_name)
            style_info = self._extract_para_style(para)

            blocks.append({
                'id': f'block_{block_idx:04d}',
                'type': block_type,
                'text': text,
                'originalText': text,
                'position': {'x': 0, 'y': 0, 'width': 0, 'height': 0},
                'style': style_info,
            })
            block_idx += 1

        return blocks

    def _extract_tables(self, doc: Document) -> list:
        """Extract tables with merged cell detection."""
        tables = []
        xml_tables = doc.element.body.findall(qn('w:tbl'))

        for table_idx, (table, xml_table) in enumerate(zip(doc.tables, xml_tables)):
            merge_map = self._build_merge_map(xml_table, table)
            cells = self._extract_table_cells(table, merge_map)

            tables.append({
                'id': f'table_{table_idx:04d}',
                'rows': len(table.rows),
                'cols': len(table.columns) if table.rows else 0,
                'cells': cells,
                'position': {'x': 0, 'y': 0, 'width': 0, 'height': 0},
            })

        return tables

    def _build_merge_map(self, xml_table, table) -> dict:
        """Build a map of (row, col) -> {'rowSpan': n, 'colSpan': n, 'isMerged': bool}.

        Detects both horizontal (gridSpan) and vertical (vMerge) merges.
        """
        merge_map = {}
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

        xml_rows = xml_table.findall(qn('w:tr'))
        for row_idx, xml_row in enumerate(xml_rows):
            col_idx = 0
            for tc in xml_row.findall(qn('w:tc')):
                # Horizontal merge: gridSpan attribute
                tc_pr = tc.find(qn('w:tcPr'))
                grid_span = 1
                if tc_pr is not None:
                    gs = tc_pr.find(qn('w:gridSpan'))
                    if gs is not None:
                        try:
                            grid_span = int(gs.get(qn('w:val'), '1'))
                        except (ValueError, TypeError):
                            grid_span = 1

                # Vertical merge: vMerge attribute
                v_merge = 'none'
                if tc_pr is not None:
                    vm = tc_pr.find(qn('w:vMerge'))
                    if vm is not None:
                        v_merge = vm.get(qn('w:val'), 'continue')
                        if v_merge is None:
                            v_merge = 'continue'

                merge_info = {
                    'colSpan': grid_span,
                    'rowSpan': 1,
                    'isMerged': grid_span > 1 or v_merge != 'none',
                    'vMerge': v_merge,
                }

                if v_merge == 'restart':
                    # Count how many rows this vertical merge spans
                    merge_info['rowSpan'] = self._count_vmerge_rows(xml_rows, row_idx, col_idx)

                merge_map[(row_idx, col_idx)] = merge_info

                # Mark cells that are continuations of horizontal merges
                for offset in range(1, grid_span):
                    merge_map[(row_idx, col_idx + offset)] = {
                        'colSpan': 0, 'rowSpan': 0,
                        'isMerged': True, 'vMerge': 'continue',
                    }

                col_idx += grid_span

        return merge_map

    def _count_vmerge_rows(self, xml_rows, start_row, col_idx) -> int:
        """Count how many consecutive rows are vertically merged starting from start_row."""
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        count = 1
        for row_idx in range(start_row + 1, len(xml_rows)):
            tcs = xml_rows[row_idx].findall(qn('w:tc'))
            if col_idx >= len(tcs):
                break
            tc = tcs[col_idx]
            tc_pr = tc.find(qn('w:tcPr'))
            if tc_pr is None:
                break
            vm = tc_pr.find(qn('w:vMerge'))
            if vm is None:
                break
            val = vm.get(qn('w:val'), 'continue')
            if val == 'restart':
                break
            count += 1
        return count

    def _extract_table_cells(self, table, merge_map: dict) -> list:
        """Extract cell text with merge info. Skips continuation cells.

        Uses cell._tc identity to deduplicate merged cells — python-docx's
        row.cells expands merged cells so the same _tc appears multiple times.
        """
        cells = []
        processed_tcs = set()  # Track underlying XML elements to deduplicate

        for row_idx, row in enumerate(table.rows):
            row_cells = []
            for col_idx, cell in enumerate(row.cells):
                # Core fix: skip already-processed merged cells
                cell_tc = cell._tc
                if cell_tc in processed_tcs:
                    continue
                processed_tcs.add(cell_tc)

                info = merge_map.get((row_idx, col_idx), {})

                # Also skip cells that are continuations of a merge (colSpan=0)
                if info.get('colSpan', 1) == 0:
                    continue

                text = _clean_text(cell.text)
                row_cells.append({
                    'text': text,
                    'originalText': text,
                    'rowSpan': info.get('rowSpan', 1),
                    'colSpan': info.get('colSpan', 1),
                    'isMerged': info.get('isMerged', False),
                })
            cells.append(row_cells)

        return cells

    @staticmethod
    def _detect_type(style_name: str) -> str:
        lower = style_name.lower()
        if 'heading' in lower or 'title' in lower:
            return 'heading'
        if 'list' in lower:
            return 'list'
        if 'caption' in lower:
            return 'caption'
        return 'paragraph'

    @staticmethod
    def _extract_style(para) -> dict:
        """Extract font style from the first run of a paragraph."""
        runs = para.runs
        font_name = ''
        font_size = 12
        bold = False
        italic = False

        if runs:
            run = runs[0]
            if run.font.name:
                font_name = run.font.name
            if run.font.size:
                font_size = run.font.size.pt
            bold = bool(run.bold)
            italic = bool(run.italic)

        return {
            'fontFamily': font_name or 'Calibri',
            'fontSize': font_size,
            'bold': bold,
            'italic': italic,
        }

    def _extract_para_style(self, para) -> dict:
        """Extract paragraph style — delegates to static helper for backwards compat."""
        return self._extract_style(para)
